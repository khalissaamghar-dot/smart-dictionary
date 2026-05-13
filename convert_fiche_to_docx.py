import pathlib
import sys
from docx import Document
from docx.shared import Pt
import markdown2

def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(text)
    if style:
        p.style = style
    return p

def markdown_to_docx(md_path, docx_path):
    doc = Document()
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()
    for line in lines:
        stripped = line.strip('\n')
        if stripped.startswith('# '):
            title = stripped[2:].strip()
            add_paragraph(doc, title, style='Heading 1')
        elif stripped.startswith('## '):
            title = stripped[3:].strip()
            add_paragraph(doc, title, style='Heading 2')
        elif stripped.startswith('### '):
            title = stripped[4:].strip()
            add_paragraph(doc, title, style='Heading 3')
        elif stripped.startswith('- '):
            # bullet point
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped == '' :
            # empty line, add a blank paragraph
            doc.add_paragraph('')
        else:
            add_paragraph(doc, stripped)
    doc.save(docx_path)
    print(f"Converted {md_path} to {docx_path}")

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print('Usage: python convert_fiche_to_docx.py <markdown_path> <output_docx_path>')
        sys.exit(1)
    md_path = pathlib.Path(sys.argv[1])
    out_path = pathlib.Path(sys.argv[2])
    markdown_to_docx(md_path, out_path)
